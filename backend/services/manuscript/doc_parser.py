"""Manuscript document parser — Phase IX.

Supports: PDF, DOCX, LaTeX, Markdown, plain text.
Returns a ParsedDocument with full_text, sections, metadata, counts.
All format-specific imports are deferred so missing optional packages
only affect the relevant parser, not the module as a whole.
"""
from __future__ import annotations

import io
import re
import logging
from pathlib import Path

from .models import ParsedDocument, InputFormat

log = logging.getLogger("synaptiq.manuscript.parser")

ALLOWED_MIME_TYPES = {
    "application/pdf": InputFormat.PDF,
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": InputFormat.DOCX,
    "application/msword": InputFormat.DOCX,
    "application/x-latex": InputFormat.LATEX,
    "text/x-latex": InputFormat.LATEX,
    "text/markdown": InputFormat.MARKDOWN,
    "text/plain": InputFormat.TXT,
}

MAX_CONTENT_CHARS = 80_000
MIN_CONTENT_CHARS = 150


# ── Format detection ──────────────────────────────────────────────────────────

def detect_format(filename: str, mime: str) -> InputFormat:
    if mime in ALLOWED_MIME_TYPES:
        return ALLOWED_MIME_TYPES[mime]
    ext = Path(filename).suffix.lower()
    return {
        ".pdf": InputFormat.PDF,
        ".docx": InputFormat.DOCX,
        ".doc": InputFormat.DOCX,
        ".tex": InputFormat.LATEX,
        ".md": InputFormat.MARKDOWN,
        ".markdown": InputFormat.MARKDOWN,
        ".txt": InputFormat.TXT,
    }.get(ext, InputFormat.TXT)


# ── PDF ───────────────────────────────────────────────────────────────────────

def parse_pdf(data: bytes) -> ParsedDocument:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(data))
    pages: list[str] = []
    for page in reader.pages:
        try:
            txt = page.extract_text() or ""
            pages.append(txt)
        except Exception:
            pass
    full_text = "\n".join(pages)
    doc = ParsedDocument(
        full_text=full_text,
        page_count=len(reader.pages),
        input_format=InputFormat.PDF,
    )
    _enrich(doc)
    return doc


# ── DOCX ──────────────────────────────────────────────────────────────────────

def parse_docx(data: bytes) -> ParsedDocument:
    from docx import Document
    d = Document(io.BytesIO(data))
    paragraphs: list[str] = []
    for para in d.paragraphs:
        txt = para.text.strip()
        if txt:
            paragraphs.append(txt)
    # Tables
    table_texts: list[str] = []
    for tbl in d.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                table_texts.append(" | ".join(cells))
    full_text = "\n".join(paragraphs)
    if table_texts:
        full_text += "\n\n[TABLES]\n" + "\n".join(table_texts)

    doc = ParsedDocument(
        full_text=full_text,
        table_count=len(d.tables),
        input_format=InputFormat.DOCX,
    )
    _enrich(doc)
    return doc


# ── LaTeX ─────────────────────────────────────────────────────────────────────

_LATEX_CMD = re.compile(r"\\[a-zA-Z]+(?:\[[^\]]*\])?(?:\{[^}]*\})*")
_LATEX_COMMENT = re.compile(r"%.*$", re.MULTILINE)
_LATEX_ENV = re.compile(r"\\(begin|end)\{[^}]+\}")
_LATEX_MATH_INLINE = re.compile(r"\$[^$]+\$")
_LATEX_MATH_BLOCK = re.compile(r"\$\$[^$]*\$\$", re.DOTALL)


def parse_latex(text: str) -> ParsedDocument:
    # Remove comments
    cleaned = _LATEX_COMMENT.sub("", text)
    # Remove math
    cleaned = _LATEX_MATH_BLOCK.sub(" [EQUATION] ", cleaned)
    cleaned = _LATEX_MATH_INLINE.sub(" [eq] ", cleaned)
    # Remove environment markers
    cleaned = _LATEX_ENV.sub("", cleaned)
    # Remove commands but keep their text arguments
    def _expand_cmd(m: re.Match) -> str:
        raw = m.group(0)
        # Keep content of last braced arg if it looks like text
        braces = re.findall(r"\{([^}]+)\}", raw)
        if braces:
            return braces[-1]
        return ""
    cleaned = _LATEX_CMD.sub(_expand_cmd, cleaned)
    # Collapse whitespace
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned).strip()

    doc = ParsedDocument(full_text=cleaned, input_format=InputFormat.LATEX)
    _enrich(doc)
    return doc


# ── Markdown ──────────────────────────────────────────────────────────────────

_MD_CODE_BLOCK = re.compile(r"```.*?```", re.DOTALL)
_MD_INLINE_CODE = re.compile(r"`[^`]+`")
_MD_IMAGE = re.compile(r"!\[[^\]]*\]\([^)]*\)")
_MD_LINK = re.compile(r"\[([^\]]+)\]\([^)]*\)")
_MD_HEADING = re.compile(r"^#{1,6}\s+", re.MULTILINE)
_MD_BOLD = re.compile(r"\*\*([^*]+)\*\*")
_MD_ITALIC = re.compile(r"\*([^*]+)\*")
_MD_HR = re.compile(r"^[-*_]{3,}$", re.MULTILINE)


def parse_markdown(text: str) -> ParsedDocument:
    cleaned = _MD_CODE_BLOCK.sub(" [CODE] ", text)
    cleaned = _MD_INLINE_CODE.sub(" [code] ", cleaned)
    cleaned = _MD_IMAGE.sub(" [FIGURE] ", cleaned)
    cleaned = _MD_LINK.sub(r"\1", cleaned)
    cleaned = _MD_BOLD.sub(r"\1", cleaned)
    cleaned = _MD_ITALIC.sub(r"\1", cleaned)
    cleaned = _MD_HEADING.sub("", cleaned)
    cleaned = _MD_HR.sub("", cleaned)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned).strip()

    doc = ParsedDocument(full_text=cleaned, input_format=InputFormat.MARKDOWN)
    _enrich(doc)
    return doc


# ── Plain text ────────────────────────────────────────────────────────────────

def parse_txt(text: str) -> ParsedDocument:
    cleaned = re.sub(r"\n{4,}", "\n\n\n", text).strip()
    doc = ParsedDocument(full_text=cleaned, input_format=InputFormat.TXT)
    _enrich(doc)
    return doc


# ── Enrichment ────────────────────────────────────────────────────────────────

_FIG_PATTERN = re.compile(
    r"(?:figure|fig\.?|Figure|Fig\.?)\s*\d+", re.IGNORECASE
)
_TABLE_PATTERN = re.compile(r"(?:Table|TABLE)\s*\d+", re.IGNORECASE)
_REF_SECTION = re.compile(
    r"(?:^References?\s*$|^Bibliography\s*$|^Works Cited\s*$)",
    re.MULTILINE | re.IGNORECASE,
)
_YEAR_IN_REF = re.compile(r"\(\d{4}\)|\b(19|20)\d{2}\b")


def _enrich(doc: ParsedDocument) -> None:
    """Populate counts and basic metadata from full_text."""
    text = doc.full_text

    # Word count
    doc.word_count = len(text.split())

    # Figure count
    doc.figure_count = len(set(_FIG_PATTERN.findall(text)))

    # Table count — add parser-extracted tables
    doc.table_count = max(doc.table_count, len(set(_TABLE_PATTERN.findall(text))))

    # Reference count: heuristic — lines in references section
    ref_match = _REF_SECTION.search(text)
    if ref_match:
        ref_block = text[ref_match.end():]
        # Count lines that look like references (contain years)
        ref_lines = [ln for ln in ref_block.split("\n")
                     if _YEAR_IN_REF.search(ln) and len(ln.strip()) > 20]
        doc.reference_count = len(ref_lines)

    # Title heuristic: first non-empty line ≤ 200 chars
    for line in text.split("\n"):
        line = line.strip()
        if line and len(line) <= 200 and not line.startswith("["):
            doc.title = line
            break

    # Abstract heuristic
    abstract_m = re.search(
        r"(?:abstract|summary)\s*\n+([\s\S]{50,}?)(?:\n{2,}|\Z)",
        text, re.IGNORECASE
    )
    if abstract_m:
        doc.abstract = abstract_m.group(1).strip()[:1500]

    # Keywords
    kw_m = re.search(
        r"(?:keywords?|key\s+words?)\s*[:\-–]?\s*\n?(.*)",
        text, re.IGNORECASE
    )
    if kw_m:
        kw_line = kw_m.group(1).strip()
        # Split on comma, semicolon, or bullet
        kws = re.split(r"[;,•·]", kw_line)
        doc.keywords = [k.strip() for k in kws if k.strip()][:10]

    # Truncate text to max
    if len(doc.full_text) > MAX_CONTENT_CHARS:
        doc.full_text = doc.full_text[:MAX_CONTENT_CHARS]


# ── Dispatcher ────────────────────────────────────────────────────────────────

def parse_document(data: bytes | str, fmt: InputFormat) -> ParsedDocument:
    """Route to the correct parser based on format."""
    if fmt == InputFormat.PDF:
        if not isinstance(data, bytes):
            raise ValueError("PDF requires bytes input")
        return parse_pdf(data)
    if fmt == InputFormat.DOCX:
        if not isinstance(data, bytes):
            raise ValueError("DOCX requires bytes input")
        return parse_docx(data)
    # Text-based formats
    text = data if isinstance(data, str) else data.decode("utf-8", errors="replace")
    if fmt == InputFormat.LATEX:
        return parse_latex(text)
    if fmt == InputFormat.MARKDOWN:
        return parse_markdown(text)
    return parse_txt(text)
