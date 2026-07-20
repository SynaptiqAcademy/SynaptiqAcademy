"""DOCX extractor — uses python-docx if available, falls back to XML parsing."""
from __future__ import annotations

import io
import re
import zipfile

from services.knowledge.ingestion.extractors.base import DocumentExtractor, ExtractedDocument
from services.knowledge.models import DocumentMetadata

_HEADING_STYLES = {"heading 1", "heading 2", "heading 3", "title"}


class DOCXExtractor(DocumentExtractor):
    @property
    def supported_types(self) -> list[str]:
        return ["docx", "doc"]

    def extract(self, content: bytes, filename: str = "") -> ExtractedDocument:
        text, sections, method = self._try_python_docx(content)
        if not text.strip():
            text, sections, method = self._xml_fallback(content)

        doi_m = re.search(r"\b10\.\d{4,9}/\S+", text)
        year_m = re.search(r"\b(19|20)\d{2}\b", text[:2000])
        meta = DocumentMetadata(
            title=filename.replace(".docx", "").replace(".doc", ""),
            doi=doi_m.group(0) if doi_m else "",
            publication_year=int(year_m.group(0)) if year_m else None,
        )
        return ExtractedDocument(
            text=text,
            sections=sections,
            metadata=meta,
            extraction_method=method,
        )

    def _try_python_docx(self, content: bytes) -> tuple[str, list[dict], str]:
        try:
            from docx import Document
            doc = Document(io.BytesIO(content))
        except ImportError:
            return "", [], ""
        except Exception:
            return "", [], ""

        sections: list[dict] = []
        current_heading = ""
        current_paras: list[str] = []

        for para in doc.paragraphs:
            style_name = (para.style.name or "").lower()
            text = para.text.strip()
            if not text:
                continue
            if style_name in _HEADING_STYLES:
                if current_paras:
                    sections.append({
                        "heading": current_heading,
                        "text": "\n".join(current_paras),
                        "page": None,
                    })
                current_heading = text
                current_paras = []
            else:
                current_paras.append(text)

        if current_paras:
            sections.append({
                "heading": current_heading,
                "text": "\n".join(current_paras),
                "page": None,
            })

        full_text = "\n\n".join(
            f"{s['heading']}\n{s['text']}" if s["heading"] else s["text"]
            for s in sections
        )
        return full_text, sections or [{"heading": "", "text": full_text, "page": None}], "python-docx"

    def _xml_fallback(self, content: bytes) -> tuple[str, list[dict], str]:
        """Parse word/document.xml directly — no library needed."""
        try:
            with zipfile.ZipFile(io.BytesIO(content)) as z:
                if "word/document.xml" not in z.namelist():
                    return "", [], "xml-fallback-failed"
                xml = z.read("word/document.xml").decode("utf-8", errors="replace")
        except Exception:
            return "", [], "xml-fallback-failed"

        xml = re.sub(r"<w:hyperlink[^>]*>", "", xml)
        xml = re.sub(r"</w:hyperlink>", "", xml)
        runs = re.findall(r"<w:t[^>]*>([^<]*)</w:t>", xml)
        text = " ".join(runs)
        text = re.sub(r"\s+", " ", text).strip()
        return text, [{"heading": "", "text": text, "page": None}], "xml-fallback"
